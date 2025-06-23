"""
Resume Parser Module
Extracts text and information from PDF and DOCX resume files
"""

import os
import re
import logging
from typing import Dict, List, Optional, Tuple
import PyPDF2
import docx
from docx import Document

# Simple imports for basic functionality
try:
    import docx2txt
except ImportError:
    docx2txt = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import spacy
    from spacy.matcher import Matcher
except ImportError:
    spacy = None
    Matcher = None

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# === Simple Functions (as requested) ===

def extract_text_from_pdf(path):
    """Simple PDF text extraction"""
    text = ""
    if fitz:  # Use PyMuPDF if available
        doc = fitz.open(path)
        for page in doc:
            text += page.get_text()
        doc.close()
    else:  # Fallback to PyPDF2
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
    return text


def extract_text_from_docx(path):
    """Simple DOCX text extraction"""
    if docx2txt:  # Use docx2txt if available
        return docx2txt.process(path)
    else:  # Fallback to python-docx
        try:
            doc = Document(path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""


def parse_resume(path):
    """Simple resume parsing function"""
    if path.endswith(".pdf"):
        text = extract_text_from_pdf(path)
    else:
        text = extract_text_from_docx(path)

    # Extract email using regex
    email = re.findall(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    
    # Extract name (assume first non-empty line)
    lines = text.split("\n")
    name = "Candidate"
    for line in lines:
        line = line.strip()
        if line and len(line) > 2 and len(line) < 50:
            # Basic validation for name
            if not any(char.isdigit() for char in line) and '@' not in line:
                name = line
                break
    
    return text, email[0] if email else "", name


# === Advanced Class (existing functionality) ===

class ResumeParser:
    def __init__(self):
        """Initialize the resume parser with NLP model"""
        try:
            # Try to load spaCy model for NLP processing
            import spacy
            from spacy.matcher import Matcher
            self.nlp = spacy.load("en_core_web_sm")
            self.matcher = Matcher(self.nlp.vocab)
            self._setup_patterns()
            logger.info("spaCy model loaded successfully")
        except (OSError, ImportError) as e:
            logger.warning(f"spaCy model not available: {str(e)}")
            logger.info("Falling back to basic text processing without NLP features")
            self.nlp = None
            self.matcher = None
    
    def _setup_patterns(self):
        """Setup patterns for entity extraction"""
        if not self.matcher:
            return
        
        try:
            # Email pattern
            email_pattern = [{"LIKE_EMAIL": True}]
            self.matcher.add("EMAIL", [email_pattern])
            
            # Phone pattern (simplified for basic matching)
            phone_pattern = [{"TEXT": {"REGEX": r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"}}]
            self.matcher.add("PHONE", [phone_pattern])
        except Exception as e:
            logger.warning(f"Could not setup spaCy patterns: {str(e)}")
            self.matcher = None
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from resume file based on extension"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else None
    
    def extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text"""
        phone_patterns = [
            r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{10,15}',
            r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        return None
    
    def extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name from text (usually first line or header)"""
        lines = text.split('\n')
        # Look for name in first few lines
        for line in lines[:5]:
            line = line.strip()
            if line and len(line.split()) <= 4 and len(line) > 3:
                # Basic validation: should contain letters and reasonable length
                if re.match(r'^[A-Za-z\s.]+$', line) and 3 <= len(line) <= 50:
                    return line
        return None
    
    def extract_education(self, text: str) -> List[str]:
        """Extract education information from text"""
        education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'university',
            'college', 'institute', 'school', 'education', 'b.s.', 'b.a.',
            'm.s.', 'm.a.', 'ph.d.', 'mba', 'certification'
        ]
        
        education_info = []
        lines = text.lower().split('\n')
        
        for line in lines:
            if any(keyword in line for keyword in education_keywords):
                education_info.append(line.strip())
        
        return education_info
    
    def extract_experience(self, text: str) -> List[str]:
        """Extract work experience from text"""
        experience_keywords = [
            'experience', 'work', 'employment', 'career', 'position',
            'job', 'role', 'worked', 'company', 'organization'
        ]
        
        experience_info = []
        lines = text.lower().split('\n')
        
        # Look for date patterns that might indicate work periods
        date_pattern = r'\d{4}|\d{1,2}/\d{4}|\d{1,2}/\d{1,2}/\d{4}'
        
        for line in lines:
            if (any(keyword in line for keyword in experience_keywords) or
                re.search(date_pattern, line)):
                experience_info.append(line.strip())
        
        return experience_info
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        skills_section_keywords = ['skills', 'technical skills', 'competencies', 'technologies']
        skills = []
        
        lines = text.split('\n')
        in_skills_section = False
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if we're entering a skills section
            if any(keyword in line_lower for keyword in skills_section_keywords):
                in_skills_section = True
                continue
            
            # Check if we're leaving skills section (new major section)
            if in_skills_section and line_lower in ['experience', 'education', 'work history', 'employment']:
                in_skills_section = False
            
            # Extract skills if we're in skills section
            if in_skills_section and line.strip():
                # Split by common delimiters
                skill_items = re.split(r'[,•·\-\|\n]', line)
                for skill in skill_items:
                    skill = skill.strip()
                    if skill and len(skill) > 1:
                        skills.append(skill)
        
        return skills
    
    def parse_resume(self, file_path: str) -> Dict:
        """
        Parse resume and extract all relevant information
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing extracted information
        """
        try:
            # Extract text
            text = self.extract_text(file_path)
            
            if not text:
                return {
                    'success': False,
                    'error': 'Could not extract text from file'
                }
            
            # Extract information
            parsed_info = {
                'success': True,
                'file_path': file_path,
                'raw_text': text,
                'name': self.extract_name(text),
                'email': self.extract_email(text),
                'phone': self.extract_phone(text),
                'education': self.extract_education(text),
                'experience': self.extract_experience(text),
                'skills': self.extract_skills(text),
                'word_count': len(text.split()),
                'character_count': len(text)
            }
            
            logger.info(f"Successfully parsed resume: {file_path}")
            return parsed_info
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def get_resume_sections(self, text: str) -> Dict[str, str]:
        """Identify and extract different sections of the resume"""
        sections = {}
        current_section = "header"
        current_content = []
        
        common_section_headers = {
            'summary': ['summary', 'profile', 'objective', 'about'],
            'experience': ['experience', 'work history', 'employment', 'career'],
            'education': ['education', 'academic', 'qualifications'],
            'skills': ['skills', 'technical skills', 'competencies'],
            'projects': ['projects', 'portfolio'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'awards': ['awards', 'achievements', 'honors']
        }
        
        lines = text.split('\n')
        
        for line in lines:
            line_clean = line.strip().lower()
            
            # Check if this line is a section header
            section_found = False
            for section_name, keywords in common_section_headers.items():
                if any(keyword in line_clean for keyword in keywords):
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = section_name
                    current_content = []
                    section_found = True
                    break
            
            if not section_found:
                current_content.append(line)
        
        # Save the last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
